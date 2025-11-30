"""
DOCX Resume Builder

Builds ATS-optimized .docx files from structured resume data.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Any
import yaml
from pathlib import Path


class DocxBuilder:
    """Build ATS-optimized DOCX files from resume data."""
    
    def __init__(self, template_path: Path = None):
        """
        Initialize DOCX builder.
        
        Args:
            template_path: Optional path to .docx template file
        """
        self.document = Document(template_path) if template_path else Document()
        self.styles = self._load_styles()
        self._setup_document()
    
    def _load_styles(self) -> Dict:
        """Load ATS formatting styles from YAML."""
        styles_path = Path(__file__).parent / "templates" / "styles.yaml"
        
        if styles_path.exists():
            with open(styles_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        else:
            # Fallback default styles
            return self._default_styles()
    
    def _default_styles(self) -> Dict:
        """Return default ATS styles if config not found."""
        return {
            "fonts": {
                "body": {"name": "Calibri", "size": 11, "fallback": "Arial"},
                "name": {"name": "Calibri", "size": 16, "bold": True, "fallback": "Arial"},
                "section_header": {"name": "Calibri", "size": 12, "bold": True, "fallback": "Arial"},
                "job_title": {"name": "Calibri", "size": 11, "bold": True, "fallback": "Arial"},
            },
            "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            "spacing": {"line": 1.0, "paragraph": 6, "section": 12},
        }
    
    def _setup_document(self):
        """Set up document margins and default formatting."""
        # Set margins
        sections = self.document.sections
        for section in sections:
            margins = self.styles.get("margins", {})
            section.top_margin = Inches(margins.get("top", 1.0))
            section.bottom_margin = Inches(margins.get("bottom", 1.0))
            section.left_margin = Inches(margins.get("left", 1.0))
            section.right_margin = Inches(margins.get("right", 1.0))
    
    def build(self, resume_data: Dict[str, Any]) -> Document:
        """
        Build complete resume document.
        
        Args:
            resume_data: Structured resume data from parser
        
        Returns:
            Document object
        """
        # Add name and contact
        self._add_header(resume_data.get("name", ""), resume_data.get("contact", {}))
        
        # Add summary
        if resume_data.get("summary"):
            self._add_section("Summary", resume_data["summary"])
        
        # Add skills
        if resume_data.get("skills"):
            self._add_skills_section(resume_data["skills"])
        
        # Add experience
        if resume_data.get("experience"):
            self._add_experience_section(resume_data["experience"])
        
        # Add education
        if resume_data.get("education"):
            self._add_education_section(resume_data["education"])
        
        # Add certifications
        if resume_data.get("certifications"):
            self._add_certifications_section(resume_data["certifications"])
        
        # Add any other sections
        for section_name, content in resume_data.get("raw_sections", {}).items():
            self._add_section(section_name, content)
        
        return self.document
    
    def _add_header(self, name: str, contact: Dict[str, str]):
        """Add name and contact information header."""
        # Name
        name_para = self.document.add_paragraph()
        name_run = name_para.add_run(name.upper())
        
        font_config = self.styles["fonts"]["name"]
        name_run.font.name = font_config.get("name", "Calibri")
        name_run.font.size = Pt(font_config.get("size", 16))
        name_run.bold = font_config.get("bold", True)
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.space_after = Pt(6)
        
        # Contact information
        contact_lines = []
        
        # Line 1: Location, Email, Phone
        line1_parts = []
        if contact.get("location"):
            line1_parts.append(contact["location"])
        if contact.get("email"):
            line1_parts.append(contact["email"])
        if contact.get("phone"):
            line1_parts.append(contact["phone"])
        
        if line1_parts:
            contact_lines.append("  ".join(line1_parts))
        
        # Line 2: LinkedIn, GitHub, Website
        line2_parts = []
        if contact.get("linkedin"):
            line2_parts.append(f"LinkedIn: {contact['linkedin']}")
        if contact.get("github"):
            line2_parts.append(f"GitHub: {contact['github']}")
        if contact.get("website"):
            line2_parts.append(contact["website"])
        
        if line2_parts:
            contact_lines.append(" | ".join(line2_parts))
        
        # Add contact lines
        for contact_line in contact_lines:
            contact_para = self.document.add_paragraph()
            contact_run = contact_para.add_run(contact_line)
            
            font_config = self.styles["fonts"]["body"]
            contact_run.font.name = font_config.get("name", "Calibri")
            contact_run.font.size = Pt(font_config.get("size", 11))
            contact_run.font.color.rgb = RGBColor(0, 0, 0)
            
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.space_after = Pt(3)
        
        # Add space after contact
        spacing_para = self.document.add_paragraph()
        spacing_para.space_after = Pt(self.styles["spacing"]["section"])
    
    def _add_section(self, title: str, content: str):
        """Add a simple text section."""
        # Section title
        self._add_section_header(title)
        
        # Section content
        if content.strip():
            para = self.document.add_paragraph()
            run = para.add_run(content.strip())
            
            font_config = self.styles["fonts"]["body"]
            run.font.name = font_config.get("name", "Calibri")
            run.font.size = Pt(font_config.get("size", 11))
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            para.space_after = Pt(self.styles["spacing"]["section"])
    
    def _add_section_header(self, title: str):
        """Add a section header."""
        header_para = self.document.add_paragraph()
        header_run = header_para.add_run(title.upper())
        
        font_config = self.styles["fonts"]["section_header"]
        header_run.font.name = font_config.get("name", "Calibri")
        header_run.font.size = Pt(font_config.get("size", 12))
        header_run.bold = font_config.get("bold", True)
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        
        header_para.space_before = Pt(self.styles["spacing"].get("before_header", 6))
        header_para.space_after = Pt(self.styles["spacing"].get("after_header", 3))
        
        # Add bottom border to section header
        self._add_paragraph_border(header_para, bottom=True)
    
    def _add_paragraph_border(self, paragraph, bottom=True):
        """Add border to paragraph."""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        if bottom:
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')  # Border size
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), '000000')
            pBdr.append(bottom_border)
        
        pPr.append(pBdr)
    
    def _add_skills_section(self, skills: Dict[str, List[str]]):
        """Add skills section with categories."""
        self._add_section_header("Core Skills")
        
        for category, skill_list in skills.items():
            if not skill_list:
                continue
            
            para = self.document.add_paragraph()
            
            # Category name (bold)
            category_run = para.add_run(f"{category}")
            font_config = self.styles["fonts"]["body"]
            category_run.font.name = font_config.get("name", "Calibri")
            category_run.font.size = Pt(font_config.get("size", 11))
            category_run.bold = True
            category_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Skills (comma-separated)
            skills_text = ", ".join(skill_list) if isinstance(skill_list[0], str) else ", ".join([str(s) for s in skill_list])
            skills_run = para.add_run(f"\n{skills_text}")
            skills_run.font.name = font_config.get("name", "Calibri")
            skills_run.font.size = Pt(font_config.get("size", 11))
            skills_run.font.color.rgb = RGBColor(0, 0, 0)
            
            para.space_after = Pt(self.styles["spacing"]["paragraph"])
        
        # Add space after section
        spacing_para = self.document.add_paragraph()
        spacing_para.space_after = Pt(self.styles["spacing"]["section"])
    
    def _add_experience_section(self, experiences: List[Dict[str, Any]]):
        """Add professional experience section."""
        self._add_section_header("Professional Experience")
        
        for i, job in enumerate(experiences):
            # Job title and company
            job_para = self.document.add_paragraph()
            
            # Title (bold)
            title_run = job_para.add_run(job.get("title", ""))
            font_config = self.styles["fonts"]["job_title"]
            title_run.font.name = font_config.get("name", "Calibri")
            title_run.font.size = Pt(font_config.get("size", 11))
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Company and dates
            company_parts = []
            if job.get("company"):
                company_parts.append(job["company"])
            if job.get("location"):
                company_parts.append(job["location"])
            
            if company_parts:
                company_text = " | ".join(company_parts)
                company_run = job_para.add_run(f" | {company_text}")
                font_config = self.styles["fonts"]["body"]
                company_run.font.name = font_config.get("name", "Calibri")
                company_run.font.size = Pt(font_config.get("size", 11))
                company_run.font.color.rgb = RGBColor(0, 0, 0)
            
            job_para.space_after = Pt(3)
            
            # Dates (italic)
            if job.get("dates"):
                dates_para = self.document.add_paragraph()
                dates_run = dates_para.add_run(job["dates"])
                
                font_config = self.styles["fonts"].get("dates", self.styles["fonts"]["body"])
                dates_run.font.name = font_config.get("name", "Calibri")
                dates_run.font.size = Pt(font_config.get("size", 11))
                dates_run.italic = True
                dates_run.font.color.rgb = RGBColor(0, 0, 0)
                
                dates_para.space_after = Pt(6)
            
            # Description (if present)
            if job.get("description"):
                desc_para = self.document.add_paragraph()
                desc_run = desc_para.add_run(job["description"])
                
                font_config = self.styles["fonts"]["body"]
                desc_run.font.name = font_config.get("name", "Calibri")
                desc_run.font.size = Pt(font_config.get("size", 11))
                desc_run.font.color.rgb = RGBColor(0, 0, 0)
                
                desc_para.space_after = Pt(6)
            
            # Achievements (bullets)
            for achievement in job.get("achievements", []):
                bullet_para = self.document.add_paragraph(style='List Bullet')
                bullet_run = bullet_para.add_run(achievement)
                
                font_config = self.styles["fonts"]["body"]
                bullet_run.font.name = font_config.get("name", "Calibri")
                bullet_run.font.size = Pt(font_config.get("size", 11))
                bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                
                bullet_para.space_after = Pt(3)
                
                # Set bullet indent
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                bullet_para.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Add space between jobs
            if i < len(experiences) - 1:
                spacing_para = self.document.add_paragraph()
                spacing_para.space_after = Pt(self.styles["spacing"]["paragraph"])
        
        # Add space after section
        spacing_para = self.document.add_paragraph()
        spacing_para.space_after = Pt(self.styles["spacing"]["section"])
    
    def _add_education_section(self, education: List[Dict[str, str]]):
        """Add education section."""
        self._add_section_header("Education")
        
        for edu in education:
            edu_para = self.document.add_paragraph()
            
            # Degree (bold)
            degree_run = edu_para.add_run(edu.get("degree", ""))
            font_config = self.styles["fonts"]["body"]
            degree_run.font.name = font_config.get("name", "Calibri")
            degree_run.font.size = Pt(font_config.get("size", 11))
            degree_run.bold = True
            degree_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # School and year
            edu_parts = []
            if edu.get("school"):
                edu_parts.append(edu["school"])
            if edu.get("year"):
                edu_parts.append(f"({edu['year']})")
            
            if edu_parts:
                edu_text = " - ".join(edu_parts) if len(edu_parts) == 2 else edu_parts[0]
                edu_run = edu_para.add_run(f" | {edu_text}")
                edu_run.font.name = font_config.get("name", "Calibri")
                edu_run.font.size = Pt(font_config.get("size", 11))
                edu_run.font.color.rgb = RGBColor(0, 0, 0)
            
            edu_para.space_after = Pt(self.styles["spacing"]["paragraph"])
            
            # Details
            for detail in edu.get("details", []):
                if detail.strip():
                    detail_para = self.document.add_paragraph()
                    detail_run = detail_para.add_run(detail)
                    detail_run.font.name = font_config.get("name", "Calibri")
                    detail_run.font.size = Pt(font_config.get("size", 11))
                    detail_run.font.color.rgb = RGBColor(0, 0, 0)
                    detail_para.space_after = Pt(3)
        
        # Add space after section
        spacing_para = self.document.add_paragraph()
        spacing_para.space_after = Pt(self.styles["spacing"]["section"])
    
    def _add_certifications_section(self, certifications: List[str]):
        """Add certifications section."""
        self._add_section_header("Certifications")
        
        for cert in certifications:
            cert_para = self.document.add_paragraph(style='List Bullet')
            cert_run = cert_para.add_run(cert)
            
            font_config = self.styles["fonts"]["body"]
            cert_run.font.name = font_config.get("name", "Calibri")
            cert_run.font.size = Pt(font_config.get("size", 11))
            cert_run.font.color.rgb = RGBColor(0, 0, 0)
            
            cert_para.space_after = Pt(3)
        
        # Add space after section
        spacing_para = self.document.add_paragraph()
        spacing_para.space_after = Pt(self.styles["spacing"]["section"])
    
    def save(self, output_path: Path):
        """
        Save document to file.
        
        Args:
            output_path: Path where to save the .docx file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_path))

