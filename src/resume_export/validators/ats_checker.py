"""
ATS Compliance Validator

Validates resume files for ATS (Applicant Tracking System) compliance.
"""

from docx import Document
from pathlib import Path
from typing import Tuple, List, Dict
import yaml


class ValidationResult:
    """Container for validation results."""
    
    def __init__(self, check_name: str, passed: bool, message: str, severity: str = "error"):
        """
        Initialize validation result.
        
        Args:
            check_name: Name of the validation check
            passed: Whether check passed
            message: Human-readable message
            severity: 'critical', 'warning', or 'info'
        """
        self.check_name = check_name
        self.passed = passed
        self.message = message
        self.severity = severity
    
    def __str__(self):
        status = "✅" if self.passed else "❌" if self.severity == "critical" else "⚠️"
        return f"{status} {self.check_name}: {self.message}"


class ATSValidator:
    """Validate resume for ATS compliance."""
    
    # ATS-friendly fonts
    ATS_FRIENDLY_FONTS = ["Calibri", "Arial", "Times New Roman", "Georgia", "Helvetica"]
    
    # File size limit (in MB)
    MAX_FILE_SIZE_MB = 1.0
    
    # Recommended page count
    MIN_PAGES = 1
    MAX_PAGES = 2
    
    def __init__(self, styles_path: Path = None):
        """
        Initialize ATS validator.
        
        Args:
            styles_path: Optional path to styles.yaml config
        """
        self.styles = self._load_styles(styles_path)
    
    def _load_styles(self, styles_path: Path = None) -> Dict:
        """Load ATS styles configuration."""
        if styles_path is None:
            styles_path = Path(__file__).parent.parent / "templates" / "styles.yaml"
        
        if styles_path.exists():
            with open(styles_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        
        return {}
    
    def validate(self, docx_path: Path) -> Tuple[bool, List[ValidationResult]]:
        """
        Run all ATS validation checks.
        
        Args:
            docx_path: Path to .docx file to validate
        
        Returns:
            Tuple of (is_valid, list of validation results)
        """
        results = []
        
        # Check file exists
        if not Path(docx_path).exists():
            results.append(ValidationResult(
                "File Existence",
                False,
                f"File not found: {docx_path}",
                "critical"
            ))
            return False, results
        
        # Load document
        try:
            doc = Document(docx_path)
        except Exception as e:
            results.append(ValidationResult(
                "File Format",
                False,
                f"Cannot open as .docx file: {e}",
                "critical"
            ))
            return False, results
        
        # Run validation checks
        results.append(self._check_file_size(docx_path))
        results.append(self._check_fonts(doc))
        results.append(self._check_tables(doc))
        results.append(self._check_images(doc))
        results.append(self._check_sections(doc))
        results.append(self._check_page_length(doc))
        results.extend(self._check_formatting(doc))
        
        # Determine overall pass/fail
        critical_failures = [r for r in results if not r.passed and r.severity == "critical"]
        is_valid = len(critical_failures) == 0
        
        return is_valid, results
    
    def _check_file_size(self, docx_path: Path) -> ValidationResult:
        """Check if file size is under limit."""
        file_size_mb = Path(docx_path).stat().st_size / (1024 * 1024)
        
        if file_size_mb <= self.MAX_FILE_SIZE_MB:
            return ValidationResult(
                "File Size",
                True,
                f"File size: {file_size_mb:.2f}MB (under {self.MAX_FILE_SIZE_MB}MB limit)",
                "info"
            )
        else:
            return ValidationResult(
                "File Size",
                False,
                f"File size ({file_size_mb:.2f}MB) exceeds {self.MAX_FILE_SIZE_MB}MB limit",
                "critical"
            )
    
    def _check_fonts(self, doc: Document) -> ValidationResult:
        """Check if fonts are ATS-friendly."""
        non_ats_fonts = set()
        
        for para in doc.paragraphs:
            for run in para.runs:
                font_name = run.font.name
                if font_name and font_name not in self.ATS_FRIENDLY_FONTS:
                    non_ats_fonts.add(font_name)
        
        if not non_ats_fonts:
            return ValidationResult(
                "Fonts",
                True,
                "All fonts are ATS-friendly",
                "info"
            )
        else:
            fonts_list = ", ".join(non_ats_fonts)
            return ValidationResult(
                "Fonts",
                False,
                f"Non-ATS fonts detected: {fonts_list}. Use {', '.join(self.ATS_FRIENDLY_FONTS)}",
                "warning"
            )
    
    def _check_tables(self, doc: Document) -> ValidationResult:
        """Check for tables (not ATS-friendly)."""
        table_count = len(doc.tables)
        
        if table_count == 0:
            return ValidationResult(
                "Tables",
                True,
                "No tables found (good for ATS)",
                "info"
            )
        else:
            return ValidationResult(
                "Tables",
                False,
                f"Document contains {table_count} table(s). ATS systems may not parse tables correctly.",
                "critical"
            )
    
    def _check_images(self, doc: Document) -> ValidationResult:
        """Check for images (not ATS-friendly)."""
        # Check for image relationships
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref.lower():
                image_count += 1
        
        if image_count == 0:
            return ValidationResult(
                "Images",
                True,
                "No images found (good for ATS)",
                "info"
            )
        else:
            return ValidationResult(
                "Images",
                False,
                f"Document contains {image_count} image(s). ATS systems cannot parse images.",
                "critical"
            )
    
    def _check_sections(self, doc: Document) -> ValidationResult:
        """Check for standard section headers."""
        standard_sections = self.styles.get("ats_rules", {}).get("standard_sections", [
            "Summary", "Professional Summary", 
            "Core Skills", "Skills", "Technical Skills",
            "Experience", "Professional Experience", "Work Experience",
            "Education"
        ])
        
        found_sections = []
        paragraph_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        for section in standard_sections:
            for text in paragraph_text:
                if section.lower() in text.lower():
                    found_sections.append(section)
                    break
        
        # Check if at least core sections exist
        has_experience = any("experience" in s.lower() for s in found_sections)
        has_education = any("education" in s.lower() for s in found_sections)
        
        if has_experience and has_education:
            return ValidationResult(
                "Sections",
                True,
                f"Found standard sections: {', '.join(set(found_sections))}",
                "info"
            )
        else:
            missing = []
            if not has_experience:
                missing.append("Experience")
            if not has_education:
                missing.append("Education")
            
            return ValidationResult(
                "Sections",
                False,
                f"Missing standard sections: {', '.join(missing)}",
                "warning"
            )
    
    def _check_page_length(self, doc: Document) -> ValidationResult:
        """Check if resume is appropriate length."""
        # Estimate page count (rough approximation)
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        estimated_pages = max(1, paragraph_count // 30)  # Rough estimate: 30 paragraphs per page
        
        if self.MIN_PAGES <= estimated_pages <= self.MAX_PAGES:
            return ValidationResult(
                "Length",
                True,
                f"Resume length: ~{estimated_pages} page(s) (optimal)",
                "info"
            )
        elif estimated_pages < self.MIN_PAGES:
            return ValidationResult(
                "Length",
                False,
                f"Resume may be too short (~{estimated_pages} page). Consider adding more detail.",
                "info"
            )
        else:
            return ValidationResult(
                "Length",
                False,
                f"Resume may be too long (~{estimated_pages} pages). Consider condensing to {self.MAX_PAGES} pages.",
                "warning"
            )
    
    def _check_formatting(self, doc: Document) -> List[ValidationResult]:
        """Check for proper formatting."""
        results = []
        
        # Check for consistent spacing
        spacing_issues = 0
        for para in doc.paragraphs:
            # Check for excessive empty paragraphs
            if not para.text.strip():
                spacing_issues += 1
        
        if spacing_issues > 10:
            results.append(ValidationResult(
                "Spacing",
                False,
                f"Excessive empty paragraphs ({spacing_issues}). Clean up spacing for ATS.",
                "warning"
            ))
        else:
            results.append(ValidationResult(
                "Spacing",
                True,
                "Spacing looks good",
                "info"
            ))
        
        # Check for reasonable line length
        long_lines = 0
        for para in doc.paragraphs:
            if len(para.text) > 150:
                long_lines += 1
        
        if long_lines > 5:
            results.append(ValidationResult(
                "Line Length",
                False,
                f"{long_lines} paragraphs have very long lines. Consider breaking them up.",
                "info"
            ))
        
        return results
    
    def generate_report(self, results: List[ValidationResult]) -> str:
        """
        Generate human-readable validation report.
        
        Args:
            results: List of validation results
        
        Returns:
            Formatted report string
        """
        report_lines = []
        report_lines.append("\n" + "="*70)
        report_lines.append("ATS COMPLIANCE VALIDATION REPORT")
        report_lines.append("="*70 + "\n")
        
        # Group by severity
        critical = [r for r in results if not r.passed and r.severity == "critical"]
        warnings = [r for r in results if not r.passed and r.severity == "warning"]
        info = [r for r in results if not r.passed and r.severity == "info"]
        passed = [r for r in results if r.passed]
        
        # Overall status
        if not critical:
            report_lines.append("✅ OVERALL STATUS: ATS-COMPLIANT")
        else:
            report_lines.append("❌ OVERALL STATUS: NOT ATS-COMPLIANT")
        
        report_lines.append("")
        
        # Critical issues
        if critical:
            report_lines.append("🚨 CRITICAL ISSUES (Must Fix):")
            for result in critical:
                report_lines.append(f"  {result}")
            report_lines.append("")
        
        # Warnings
        if warnings:
            report_lines.append("⚠️  WARNINGS (Should Fix):")
            for result in warnings:
                report_lines.append(f"  {result}")
            report_lines.append("")
        
        # Info
        if info:
            report_lines.append("ℹ️  SUGGESTIONS:")
            for result in info:
                report_lines.append(f"  {result}")
            report_lines.append("")
        
        # Passed checks
        report_lines.append("✅ PASSED CHECKS:")
        for result in passed:
            report_lines.append(f"  {result}")
        
        report_lines.append("\n" + "="*70)
        
        return "\n".join(report_lines)


def validate_resume(docx_path: Path) -> Tuple[bool, str]:
    """
    Convenience function to validate a resume and get report.
    
    Args:
        docx_path: Path to .docx file
    
    Returns:
        Tuple of (is_valid, report_text)
    """
    validator = ATSValidator()
    is_valid, results = validator.validate(docx_path)
    report = validator.generate_report(results)
    
    return is_valid, report

